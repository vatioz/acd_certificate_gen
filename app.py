import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from io import BytesIO
from datetime import datetime
import re
import copy


def replace_placeholders(doc, replacements):
    """
    Replace placeholders in the document with actual values.
    Placeholders format: [NAME], [DOB]
    Gender-dependent placeholders: [ZÍSKAL/A], [ABSOLVOVAL/A], [DOKONČIL/A], etc.
    """
    
    def replace_in_paragraph(paragraph):
        """Replace placeholders in a paragraph, handling split runs."""
        # Check if any placeholder exists in the full paragraph text
        full_text = paragraph.text
        new_text = full_text
        
        for key, value in replacements.items():
            new_text = new_text.replace(key, value)
        
        if full_text != new_text:
            # Preserve formatting from first run if available
            if paragraph.runs:
                # Store the font properties from the first run
                first_run = paragraph.runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
                font_bold = first_run.font.bold
                font_italic = first_run.font.italic
                font_underline = first_run.font.underline
                
                # Clear all runs except the first one
                while len(paragraph.runs) > 1:
                    paragraph.runs[-1]._element.getparent().remove(paragraph.runs[-1]._element)
                
                # Update the first run's text
                paragraph.runs[0].text = new_text
                
                # Restore formatting
                if font_name:
                    paragraph.runs[0].font.name = font_name
                if font_size:
                    paragraph.runs[0].font.size = font_size
                if font_bold is not None:
                    paragraph.runs[0].font.bold = font_bold
                if font_italic is not None:
                    paragraph.runs[0].font.italic = font_italic
                if font_underline is not None:
                    paragraph.runs[0].font.underline = font_underline
            else:
                # No existing runs, just add the text
                paragraph.add_run(new_text)
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    
    # Replace in headers
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            replace_in_paragraph(paragraph)
        
        # Replace in footers
        footer = section.footer
        for paragraph in footer.paragraphs:
            replace_in_paragraph(paragraph)
    
    return doc


def generate_certificate(template_bytes, name, dob, gender='female', counter=1):
    """
    Generate a certificate from template with provided data.
    Returns a Document object (not bytes).
    
    Args:
        template_bytes: Template document bytes
        name: Person's full name
        dob: Date of birth (formatted string)
        gender: 'male' or 'female' (default: 'female')
        counter: Certificate number in the series (default: 1)
    """
    # Load template from bytes
    doc = Document(BytesIO(template_bytes))
    
    # Define gender-dependent verb endings
    if gender == 'male':
        verb_ending = ''
        past_ending = 'l'
    else:  # female
        verb_ending = 'a'
        past_ending = 'la'
    
    # Define replacements
    replacements = {
        '[NAME]': name,
        '[DOB]': dob,
        '[YEAR]': str(datetime.now().year),
        '[COUNTER]': str(counter),
        # Common Czech verbs in past tense
        '[ZISKAL/A]': f'získal{verb_ending}',
        '[ABSOLVOVAL/A]': f'absolvoval{verb_ending}',
        '[NAROZEN/A]': f'narozen{verb_ending}'        
    }
    
    # Replace placeholders
    doc = replace_placeholders(doc, replacements)
    
    return doc


def generate_multi_certificate(template_bytes, people_list, starting_counter=1):
    """
    Generate multiple certificates in a single document.
    Each certificate on a separate page.
    
    Args:
        template_bytes: Template document bytes
        people_list: List of people with name, dob, gender
        starting_counter: Starting certificate number (increments for each person)
    """
    if not people_list:
        return None
    
    # Start with the first certificate - this preserves all template styling
    first_person = people_list[0]
    final_doc = generate_certificate(
        template_bytes,
        first_person['name'],
        first_person['dob'],
        first_person.get('gender', 'female'),
        starting_counter
    )
    
    # Add remaining certificates with page breaks
    for idx, person in enumerate(people_list[1:], start=1):
        # Generate certificate with gender and incremented counter
        temp_doc = generate_certificate(
            template_bytes,
            person['name'],
            person['dob'],
            person.get('gender', 'female'),
            starting_counter + idx
        )
        
        # Add page break to the last paragraph of the previous certificate
        # Find the last paragraph in final_doc
        from docx.oxml import OxmlElement
        last_para = None
        for elem in reversed(final_doc.element.body):
            if elem.tag.endswith('p'):
                last_para = elem
                break
        
        if last_para is not None:
            # Add page break run to the end of the last paragraph
            page_break_run = OxmlElement('w:r')
            page_break_br = OxmlElement('w:br')
            page_break_br.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'page')
            page_break_run.append(page_break_br)
            last_para.append(page_break_run)
        
        # Copy all elements from temp_doc to final_doc
        for element in temp_doc.element.body:
            # Skip section properties
            if element.tag.endswith('sectPr'):
                continue
            
            # Deep copy the element - preserves all formatting
            new_element = copy.deepcopy(element)
            
            # Append to final document
            final_doc.element.body.append(new_element)
    
    # Save to BytesIO
    output = BytesIO()
    final_doc.save(output)
    output.seek(0)
    
    return output


def main():
    st.set_page_config(
        page_title="Certificate Generator",
        page_icon="📜",
        layout="centered"
    )
    
    st.title("📜 Certificate Generator")
    st.markdown("Upload a certificate template and generate personalized certificates")
    
    # Initialize session state for template
    if 'template' not in st.session_state:
        st.session_state.template = None
        st.session_state.template_name = None
    
    # Initialize session state for people list
    if 'people_list' not in st.session_state:
        st.session_state.people_list = []
    
    # Initialize session state for starting counter
    if 'starting_counter' not in st.session_state:
        st.session_state.starting_counter = 1
    
    # File uploader
    st.subheader("1. Upload Certificate Template")
    uploaded_file = st.file_uploader(
        "Choose a .docx template file",
        type=['docx'],
        help="Upload a Word document with placeholders [NAME] and [DOB]"
    )
    
    # Store template in session state
    if uploaded_file is not None:
        st.session_state.template = uploaded_file.read()
        st.session_state.template_name = uploaded_file.name
        st.success(f"✅ Template '{uploaded_file.name}' loaded successfully!")
    
    # Show form only if template is loaded
    if st.session_state.template is not None:
        st.subheader("2. Add Certificate Holders")
        
        # Test button to add 8 sample people
        if st.button("🧪 Add 8 Test People", type="secondary"):
            st.session_state.people_list = [
                {'name': 'Jan Novák', 'dob': '01. 01. 1990', 'gender': 'male'},
                {'name': 'Petr Dvořák', 'dob': '15. 03. 1985', 'gender': 'male'},
                {'name': 'Marie Svobodová', 'dob': '22. 07. 1992', 'gender': 'female'},
                {'name': 'Eva Černá', 'dob': '10. 12. 1988', 'gender': 'female'},
                {'name': 'Tomáš Procházka', 'dob': '05. 05. 1995', 'gender': 'male'},
                {'name': 'Jana Kučerová', 'dob': '18. 09. 1991', 'gender': 'female'},
                {'name': 'Pavel Horák', 'dob': '28. 02. 1987', 'gender': 'male'},
                {'name': 'Lucie Málková', 'dob': '14. 11. 1993', 'gender': 'female'}
            ]
            st.rerun()
        
        with st.form("certificate_form"):
            col1, col2, col3 = st.columns([3, 3, 2])
            
            with col1:
                name = st.text_input(
                    "Full Name",
                    placeholder="Jan Novák",
                    help="Enter the full name (Surname Name format for filename)"
                )
            
            with col2:
                dob = st.date_input(
                    "Date of Birth",
                    value=None,
                    min_value=datetime(1900, 1, 1),
                    max_value=datetime.today(),
                    format="DD.MM.YYYY",
                    help="Select date of birth"
                )
            
            with col3:
                gender = st.radio(
                    "Gender",
                    options=['female', 'male'],
                    format_func=lambda x: '👩 Žena' if x == 'female' else '👨 Muž',
                    index=0,
                    help="Default: Žena"
                )
            
            col_a, col_b = st.columns(2)
            with col_a:
                add_person = st.form_submit_button("➕ Add Person", type="secondary")
            with col_b:
                clear_all = st.form_submit_button("🗑️ Clear All", type="secondary")
            
            if add_person:
                if not name or not dob:
                    st.error("⚠️ Please fill in all fields!")
                else:
                    # Format date in Czech format
                    formatted_dob = dob.strftime("%d. %m. %Y")
                    
                    # Add to list
                    st.session_state.people_list.append({
                        'name': name,
                        'dob': formatted_dob,
                        'gender': gender
                    })
                    gender_label = '👩' if gender == 'female' else '👨'
                    st.success(f"✅ Added {gender_label} {name}")
                    st.rerun()
            
            if clear_all:
                st.session_state.people_list = []
                st.rerun()
        
        # Display current list
        if st.session_state.people_list:
            st.subheader("3. Certificate Holders List")
            
            # Display as table
            for idx, person in enumerate(st.session_state.people_list):
                col1, col2, col3, col4 = st.columns([3, 3, 1.5, 1])
                with col1:
                    st.text(person['name'])
                with col2:
                    st.text(person['dob'])
                with col3:
                    # Gender toggle button
                    current_gender = person.get('gender', 'female')
                    gender_icon = '👩 Žena' if current_gender == 'female' else '👨 Muž'
                    if st.button(gender_icon, key=f"gender_{idx}", help="Klikněte pro změnu pohlaví"):
                        # Toggle gender
                        new_gender = 'male' if current_gender == 'female' else 'female'
                        st.session_state.people_list[idx]['gender'] = new_gender
                        st.rerun()
                with col4:
                    if st.button("🗑️", key=f"del_{idx}", help="Remove"):
                        st.session_state.people_list.pop(idx)
                        st.rerun()
            
            st.divider()
            
            # Generate all certificates button
            st.subheader("4. Generate Certificates")
            
            # Counter input
            col_counter, col_year = st.columns([2, 2])
            with col_counter:
                starting_counter = st.number_input(
                    "Starting Certificate Number",
                    min_value=1,
                    value=st.session_state.starting_counter,
                    step=1,
                    key='counter_input',
                    help=f"First certificate will be numbered with this value, then increments for each person"
                )
                st.session_state.starting_counter = starting_counter
            with col_year:
                current_year = datetime.now().year
                st.info(f"📅 Year: {current_year}")
            
            if st.button("📄 Generate All Certificates", type="primary", use_container_width=True):
                try:
                    # Generate multi-certificate document with counter
                    certificate_bytes = generate_multi_certificate(
                        st.session_state.template,
                        st.session_state.people_list,
                        starting_counter
                    )
                    
                    # Generate filename
                    if len(st.session_state.people_list) == 1:
                        filename = f"certificate_{st.session_state.people_list[0]['name'].replace(' ', '_')}.docx"
                    else:
                        filename = f"certificates_batch_{len(st.session_state.people_list)}.docx"
                    
                    # Store in session state for download
                    st.session_state.generated_cert = certificate_bytes.getvalue()
                    st.session_state.cert_filename = filename
                    
                    st.success(f"✅ Generated {len(st.session_state.people_list)} certificate(s) successfully!")
                    
                except Exception as e:
                    st.error(f"❌ Error generating certificates: {str(e)}")
        
        # Download button (outside form to avoid re-generation)
        if 'generated_cert' in st.session_state and st.session_state.people_list:
            st.subheader("5. Download Certificates")
            st.download_button(
                label=f"📥 Download {len(st.session_state.people_list)} Certificate(s)",
                data=st.session_state.generated_cert,
                file_name=st.session_state.cert_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="secondary",
                use_container_width=True
            )
    else:
        st.info("👆 Please upload a certificate template to begin.")
    
    # Sidebar with instructions
    with st.sidebar:
        st.header("ℹ️ Instructions")
        st.markdown("""
        1. **Prepare your template**: Create a Word document (.docx) with placeholders:
           - `[NAME]` - for the certificate holder's name
           - `[DOB]` - for the date of birth
           - `[COUNTER]` - certificate number (auto-increments)
           - `[YEAR]` - current year (automatic)
           - Gender-dependent verbs (Czech):
             - `[ZÍSKAL/A]` → získal / získala
             - `[ABSOLVOVAL/A]` → absolvoval / absolvovala
             - `[DOKONČIL/A]` → dokončil / dokončila
             - `[SPLNIL/A]`, `[SLOŽIL/A]`, `[VYKONAL/A]`, etc.
        
        2. **Upload** the template using the file uploader
        
        3. **Add people** - Enter name, date of birth, and gender
           - Default gender: 👩 Žena
           - Click gender button in list to toggle
           - Use test data button for quick testing
        
        4. **Review** the list and adjust gender if needed
        
        5. **Set starting number** for the certificate series
           - Counter auto-increments for each person
        
        6. **Generate** all certificates in a single document
           - Each certificate will be on a separate page
        
        7. **Download** the single file with all certificates
        
        8. **Repeat** - The template is preserved for future use!
        """)
        
        st.divider()
        st.caption("Version 0.4 - Counter & Year support")        

if __name__ == "__main__":
    main()
