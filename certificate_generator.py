"""
Certificate Generator Module

Handles all document processing logic for certificate generation.
Includes placeholder replacement, single certificate generation, and batch certificate generation.
"""

from docx import Document
from io import BytesIO
from datetime import datetime
import copy
from docx.oxml import OxmlElement


def replace_placeholders(doc, replacements):
    """
    Replace placeholders in the document with actual values.
    Placeholders format: [NAME], [DOB]
    Gender-dependent placeholders: [ZÍSKAL/A], [ABSOLVOVAL/A], [DOKONČIL/A], etc.
    """
    
    def replace_in_paragraph(paragraph):
        """Replace placeholders in a paragraph, handling split runs."""
        # Check if any placeholder exists in the full paragraph text
        full_text = paragraph.text
        new_text = full_text
        
        for key, value in replacements.items():
            new_text = new_text.replace(key, value)
        
        if full_text != new_text:
            # Preserve formatting from first run if available
            if paragraph.runs:
                # Store the font properties from the first run
                first_run = paragraph.runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
                font_bold = first_run.font.bold
                font_italic = first_run.font.italic
                font_underline = first_run.font.underline
                
                # Clear all runs except the first one
                while len(paragraph.runs) > 1:
                    paragraph.runs[-1]._element.getparent().remove(paragraph.runs[-1]._element)
                
                # Update the first run's text
                paragraph.runs[0].text = new_text
                
                # Restore formatting
                if font_name:
                    paragraph.runs[0].font.name = font_name
                if font_size:
                    paragraph.runs[0].font.size = font_size
                if font_bold is not None:
                    paragraph.runs[0].font.bold = font_bold
                if font_italic is not None:
                    paragraph.runs[0].font.italic = font_italic
                if font_underline is not None:
                    paragraph.runs[0].font.underline = font_underline
            else:
                # No existing runs, just add the text
                paragraph.add_run(new_text)
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    
    # Replace in headers
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            replace_in_paragraph(paragraph)
        
        # Replace in footers
        footer = section.footer
        for paragraph in footer.paragraphs:
            replace_in_paragraph(paragraph)
    
    return doc


def generate_certificate(template_bytes, name, dob, gender='female', counter=1):
    """
    Generate a certificate from template with provided data.
    Returns a Document object (not bytes).
    
    Args:
        template_bytes: Template document bytes
        name: Person's full name
        dob: Date of birth (formatted string)
        gender: 'male' or 'female' (default: 'female')
        counter: Certificate number in the series (default: 1)
    """
    # Load template from bytes
    doc = Document(BytesIO(template_bytes))
    
    # Define gender-dependent verb endings
    if gender == 'male':
        verb_ending = ''
        past_ending = 'l'
    else:  # female
        verb_ending = 'a'
        past_ending = 'la'
    
    # Define replacements
    replacements = {
        '[NAME]': name,
        '[DOB]': dob,
        '[YEAR]': str(year if year is not None else datetime.now().year),
        '[COUNTER]': str(counter),
        # Common Czech verbs in past tense
        '[ZISKAL/A]': f'získal{verb_ending}',
        '[ABSOLVOVAL/A]': f'absolvoval{verb_ending}',
        '[NAROZEN/A]': f'narozen{verb_ending}'        
    }
    
    # Replace placeholders
    doc = replace_placeholders(doc, replacements)
    
    return doc


def generate_multi_certificate(template_bytes, people_list, starting_counter=1, year=None):
    """
    Generate multiple certificates in a single document.
    Each certificate on a separate page.
    
    Args:
        template_bytes: Template document bytes
        people_list: List of people with name, dob, gender
        starting_counter: Starting certificate number (increments for each person)
        year: Certificate year (defaults to current year)
    """
    if not people_list:
        return None
    
    # Start with the first certificate - this preserves all template styling
    first_person = people_list[0]
    final_doc = generate_certificate(
        template_bytes,
        first_person['name'],
        first_person['dob'],
        first_person.get('gender', 'female'),
        starting_counter,
        year
    )
    
    # Add remaining certificates with page breaks
    for idx, person in enumerate(people_list[1:], start=1):
        # Generate certificate with gender and incremented counter
        temp_doc = generate_certificate(
            template_bytes,
            person['name'],
            person['dob'],
            person.get('gender', 'female'),
            starting_counter + idx,
            year
        )
        
        # Add page break to the last paragraph of the previous certificate
        # Find the last paragraph in final_doc
        last_para = None
        for elem in reversed(final_doc.element.body):
            if elem.tag.endswith('p'):
                last_para = elem
                break
        
        if last_para is not None:
            # Add page break run to the end of the last paragraph
            page_break_run = OxmlElement('w:r')
            page_break_br = OxmlElement('w:br')
            page_break_br.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'page')
            page_break_run.append(page_break_br)
            last_para.append(page_break_run)
        
        # Copy all elements from temp_doc to final_doc
        for element in temp_doc.element.body:
            # Skip section properties
            if element.tag.endswith('sectPr'):
                continue
            
            # Deep copy the element - preserves all formatting
            new_element = copy.deepcopy(element)
            
            # Append to final document
            final_doc.element.body.append(new_element)
    
    # Save to BytesIO
    output = BytesIO()
    final_doc.save(output)
    output.seek(0)
    
    return output
